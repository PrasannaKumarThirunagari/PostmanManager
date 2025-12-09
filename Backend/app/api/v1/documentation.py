"""
Documentation Export API endpoints.
Provides functionality to export Postman collections to Microsoft Word format.
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from pathlib import Path
import json
import re
from datetime import datetime
from app.config import settings

router = APIRouter()

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentationRequest(BaseModel):
    """Documentation export request model."""
    collection_id: str
    include_xss: bool = True
    include_sql: bool = True
    include_html: bool = True
    include_requests: bool = True
    include_responses: bool = True
    include_headers: bool = True
    include_body: bool = True


def sanitize_filename(filename: str) -> str:
    """Sanitize filename for filesystem compatibility."""
    # Remove or replace invalid characters
    filename = re.sub(r'[<>:"/\\|?*]', '-', filename)
    # Remove leading/trailing spaces and dots
    filename = filename.strip('. ')
    # Replace multiple spaces/hyphens with single hyphen
    filename = re.sub(r'[\s\-]+', '-', filename)
    return filename


def create_word_documentation(
    collection_path: Path,
    output_path: Path,
    options: DocumentationRequest
) -> Path:
    """Create Word documentation from Postman collection."""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="python-docx library is not installed. Please install it: pip install python-docx"
        )
    
    # Load collection
    with open(collection_path, 'r', encoding='utf-8') as f:
        collection = json.load(f)
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading(collection.get('info', {}).get('name', 'API Documentation'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Collection info
    info = collection.get('info', {})
    if info.get('description'):
        desc_para = doc.add_paragraph(info.get('description'))
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents placeholder
    doc.add_heading('Table of Contents', 1)
    toc_para = doc.add_paragraph('(Generated automatically)')
    toc_para.italic = True
    doc.add_page_break()
    
    # Process items recursively
    def process_item(item: Dict[str, Any], level: int = 1):
        """Process collection item (request or folder)."""
        if 'item' in item:
            # It's a folder
            folder_name = item.get('name', 'Unnamed Folder')
            doc.add_heading(folder_name, level)
            
            if item.get('description'):
                doc.add_paragraph(item.get('description'))
            
            doc.add_paragraph()  # Spacing
            
            # Process folder items
            for sub_item in item.get('item', []):
                process_item(sub_item, level + 1)
        else:
            # It's a request
            request_name = item.get('name', 'Unnamed Request')
            doc.add_heading(request_name, level)
            
            request_data = item.get('request', {})
            
            # Method and URL
            method = request_data.get('method', 'GET')
            url_data = request_data.get('url', {})
            url_raw = url_data.get('raw', '')
            
            method_para = doc.add_paragraph()
            method_run = method_para.add_run(f"{method} ")
            method_run.bold = True
            method_run.font.size = Pt(12)
            url_run = method_para.add_run(url_raw)
            url_run.font.size = Pt(12)
            
            doc.add_paragraph()  # Spacing
            
            # Description
            if request_data.get('description'):
                doc.add_paragraph(request_data.get('description'))
                doc.add_paragraph()
            
            # Headers
            if options.include_headers and request_data.get('header'):
                doc.add_heading('Headers', level + 1)
                headers_table = doc.add_table(rows=1, cols=2)
                headers_table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = headers_table.rows[0].cells
                header_cells[0].text = 'Key'
                header_cells[0].paragraphs[0].runs[0].bold = True
                header_cells[1].text = 'Value'
                header_cells[1].paragraphs[0].runs[0].bold = True
                
                # Data rows
                for header in request_data.get('header', []):
                    row_cells = headers_table.add_row().cells
                    row_cells[0].text = header.get('key', '')
                    row_cells[1].text = header.get('value', '')
                
                doc.add_paragraph()
            
            # Request Body
            if options.include_body and request_data.get('body'):
                doc.add_heading('Request Body', level + 1)
                body = request_data.get('body', {})
                
                if body.get('mode') == 'raw' and body.get('raw'):
                    body_para = doc.add_paragraph(body.get('raw'))
                    body_para.style = 'No Spacing'
                    # Try to format as code
                    for run in body_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                elif body.get('mode') == 'urlencoded' and body.get('urlencoded'):
                    body_table = doc.add_table(rows=1, cols=2)
                    body_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = body_table.rows[0].cells
                    header_cells[0].text = 'Key'
                    header_cells[0].paragraphs[0].runs[0].bold = True
                    header_cells[1].text = 'Value'
                    header_cells[1].paragraphs[0].runs[0].bold = True
                    
                    # Data rows
                    for param in body.get('urlencoded', []):
                        row_cells = body_table.add_row().cells
                        row_cells[0].text = param.get('key', '')
                        row_cells[1].text = param.get('value', '')
                
                doc.add_paragraph()
            
            # Responses
            if options.include_responses and item.get('response'):
                doc.add_heading('Responses', level + 1)
                
                for response in item.get('response', []):
                    status_code = response.get('code', response.get('status', '200'))
                    status_name = response.get('name', f'{status_code} Response')
                    
                    doc.add_heading(f'{status_code} - {status_name}', level + 2)
                    
                    # Response headers
                    if response.get('header'):
                        doc.add_heading('Response Headers', level + 3)
                        headers_table = doc.add_table(rows=1, cols=2)
                        headers_table.style = 'Light Grid Accent 1'
                        
                        # Header row
                        header_cells = headers_table.rows[0].cells
                        header_cells[0].text = 'Key'
                        header_cells[0].paragraphs[0].runs[0].bold = True
                        header_cells[1].text = 'Value'
                        header_cells[1].paragraphs[0].runs[0].bold = True
                        
                        # Data rows
                        for header in response.get('header', []):
                            row_cells = headers_table.add_row().cells
                            row_cells[0].text = header.get('key', '')
                            row_cells[1].text = header.get('value', '')
                        
                        doc.add_paragraph()
                    
                    # Response body
                    if response.get('body'):
                        doc.add_heading('Response Body', level + 3)
                        try:
                            body_json = json.loads(response.get('body', '{}'))
                            body_str = json.dumps(body_json, indent=2)
                        except:
                            body_str = response.get('body', '')
                        
                        body_para = doc.add_paragraph(body_str)
                        body_para.style = 'No Spacing'
                        for run in body_para.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                        
                        doc.add_paragraph()
            
            doc.add_paragraph()  # Spacing between requests
    
    # Process all items in collection
    items = collection.get('item', [])
    for item in items:
        process_item(item)
    
    # Save document
    doc.save(output_path)
    return output_path


@router.post("/export")
async def export_documentation(request: DocumentationRequest):
    """Export Postman collection to Word documentation."""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="python-docx library is not installed. Please install it: pip install python-docx"
        )
    
    # Find collection file
    collections_dir = Path(settings.postman_collections_dir)
    collection_path = None
    
    # Try to find collection file
    for api_folder in collections_dir.iterdir():
        if api_folder.is_dir():
            collection_file = api_folder / f"{request.collection_id}.postman_collection.json"
            if collection_file.exists():
                collection_path = collection_file
                break
    
    if not collection_path:
        # Try direct path
        potential_path = collections_dir / request.collection_id / f"{request.collection_id}.postman_collection.json"
        if potential_path.exists():
            collection_path = potential_path
    
    if not collection_path:
        raise HTTPException(status_code=404, detail="Collection not found")
    
    # Load collection to get API name
    with open(collection_path, 'r', encoding='utf-8') as f:
        collection = json.load(f)
    
    api_name = collection.get('info', {}).get('name', request.collection_id)
    api_name = sanitize_filename(api_name)
    
    # Create output directory (same folder as collection)
    output_dir = collection_path.parent
    output_filename = f"{api_name}_Documentation.docx"
    output_path = output_dir / output_filename
    
    # Generate documentation
    try:
        create_word_documentation(collection_path, output_path, request)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate documentation: {str(e)}")
    
    # Return file
    if output_path.exists():
        headers = {
            'Content-Disposition': f'attachment; filename="{output_filename}"'
        }
        return FileResponse(
            path=str(output_path),
            filename=output_filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers=headers
        )
    else:
        raise HTTPException(status_code=500, detail="Documentation file was not created")


@router.get("/collections")
async def get_collections_for_documentation():
    """Get list of collections available for documentation export."""
    collections_dir = Path(settings.postman_collections_dir)
    collections = []
    
    if not collections_dir.exists():
        return {"collections": []}
    
    for api_folder in collections_dir.iterdir():
        if api_folder.is_dir():
            collection_file = api_folder / f"{api_folder.name}.postman_collection.json"
            if collection_file.exists():
                try:
                    with open(collection_file, 'r', encoding='utf-8') as f:
                        collection = json.load(f)
                    
                    collections.append({
                        "id": api_folder.name,
                        "name": collection.get('info', {}).get('name', api_folder.name),
                        "description": collection.get('info', {}).get('description', '')
                    })
                except Exception as e:
                    # Skip collections that can't be read
                    continue
    
    return {"collections": collections}

